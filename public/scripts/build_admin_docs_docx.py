#!/usr/bin/env python3
"""Generate single Word file: Dainely Admin Panel Technical Documentation."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

OUT = Path(__file__).resolve().parents[1] / "Dainely-Admin-Panel-Technical-Documentation.docx"


def font(run, name="Calibri", size=11, bold=False, italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def h(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return p


def p(doc, text, bold=False, italic=False, size=11):
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(6)
    run = para.add_run(text)
    font(run, size=size, bold=bold, italic=italic)
    return para


def bullets(doc, items):
    for item in items:
        para = doc.add_paragraph(style="List Bullet")
        run = para.add_run(item)
        font(run, size=10)


def numbered(doc, items):
    for item in items:
        para = doc.add_paragraph(style="List Number")
        run = para.add_run(item)
        font(run, size=10)


def table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, header in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(header)
        font(run, size=9, bold=True)
        tcPr = cell._tc.get_or_add_tcPr()
        shd = tcPr.makeelement(
            qn("w:shd"),
            {qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): "E8EEF4"},
        )
        tcPr.append(shd)
    for ri, row in enumerate(rows):
        for ci in range(len(headers)):
            val = row[ci] if ci < len(row) else ""
            cell = t.rows[ri + 1].cells[ci]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            font(run, size=8)
    doc.add_paragraph()


def code(doc, text):
    for line in text.splitlines() or [""]:
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.left_indent = Inches(0.2)
        run = para.add_run(line if line else " ")
        font(run, name="Consolas", size=8)
        shd = para._element.get_or_add_pPr()
        shading = shd.makeelement(
            qn("w:shd"),
            {qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): "F5F5F5"},
        )
        shd.append(shading)
    doc.add_paragraph()


def build():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)

    # Cover
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(60)
    r = title.add_run("Dainely.com")
    font(r, size=28, bold=True)
    r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Admin Panel — Complete Technical Documentation")
    font(r, size=16)
    r.font.color.rgb = RGBColor(0x44, 0x44, 0x55)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_before = Pt(20)
    r = meta.add_run(
        "Admin URL: /dainely-admin-panel\n"
        "Derived from source code only — no invented features\n"
        "Audience: Developers · QA · Product Owners · AI Assistants"
    )
    font(r, size=11)

    toc = doc.add_paragraph()
    toc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    toc.paragraph_format.space_before = Pt(28)
    r = toc.add_run(
        "Contents: Overview · Architecture · Auth · Database · Features · Routes · "
        "Controllers · Models · Validation · Middleware · Jobs · Security · QA · "
        "Troubleshooting · Deployment · Technical Debt · Onboarding"
    )
    font(r, size=9, italic=True)
    doc.add_page_break()

    # 1. Executive Overview
    h(doc, "1. Executive Overview", 1)
    p(doc, "The Dainely Admin Panel is a lightweight CMS for managing Supabase-backed storefront content overlays. It does not manage Shopify product creation, payments, orders, or customer accounts.")
    p(doc, "Administrators can edit: product locale overlays and page blocks; landing pages; education page blocks; bundles; FAQs; knowledge signals; internal content links; webhook logs.")
    p(doc, "Products are synced from Shopify (webhooks/sync jobs). Admin only edits CMS overlays stored in Supabase PostgreSQL.")

    h(doc, "1.1 Technology Stack", 2)
    table(
        doc,
        ["Layer", "Choice", "Source"],
        [
            ["PHP", "^8.1", "composer.json"],
            ["Laravel", "^10.10", "composer.json"],
            ["Views", "Blade", "resources/views/admin/"],
            ["JS", "Alpine.js ^3.15", "package.json"],
            ["CSS", "Tailwind CSS ^3.4", "package.json"],
            ["Build", "Vite ^5 + laravel-vite-plugin", "package.json"],
            ["Admin DB", "Supabase PostgreSQL (connection: supabase)", "config/database.php"],
            ["Auth", "Session flag admin_authenticated", "AdminAuthController"],
            ["Permissions package", "spatie/laravel-permission installed but UNUSED", "composer.json"],
        ],
    )

    h(doc, "1.2 What Admin Is Not", 2)
    bullets(
        doc,
        [
            "No storefront customer UI (except overlays affecting PDP/meta)",
            "No payment, Stripe, Square, or order management in Admin controllers",
            "No role/permission matrix (any logged-in admin has full CMS access)",
            "No password reset / MFA (ADMIN_MFA_ENABLED exists in env.example but is not enforced)",
            "No product create/delete in admin — products come from Shopify sync",
        ],
    )
    doc.add_page_break()

    # 2. Architecture
    h(doc, "2. System Architecture", 1)
    h(doc, "2.1 Request Flow", 2)
    code(
        doc,
        "Browser\n"
        "  → Route prefix: dainely-admin-panel  (middleware: web)\n"
        "    → Session + CSRF + TrackPageViews (admin paths excluded from analytics)\n"
        "    → [protected] AdminAuth middleware (session admin_authenticated)\n"
        "      → Admin*Controller\n"
        "        → SupabaseDb::run() / ping() / flashIfSupabaseOffline()\n"
        "          → Eloquent (connection=supabase) OR ContentCatalog (virtual)\n"
        "            → Blade view extends layouts/admin.blade.php",
    )

    h(doc, "2.2 Architecture Diagram (textual Mermaid)", 2)
    code(
        doc,
        "flowchart TD\n"
        "  A[Browser] --> B[routes/web.php dainely-admin-panel]\n"
        "  B --> C{admin.auth?}\n"
        "  C -->|No| D[AdminAuthController]\n"
        "  C -->|Yes| E[Admin Module Controllers]\n"
        "  E --> F[SupabaseDb]\n"
        "  F --> G[(Supabase PostgreSQL)]\n"
        "  E --> H[ContentCatalog virtual education/blog]\n"
        "  E --> I[Blade admin views]\n"
        "  I --> J[layouts/admin Alpine + Tailwind + Vite]",
    )

    h(doc, "2.3 Design Characteristics", 2)
    bullets(
        doc,
        [
            "No repository layer — controllers use Eloquent inside SupabaseDb::run()",
            "No Form Request classes — inline $request->validate()",
            "POST-based mutations (no REST PUT/DELETE); paths use /update and /delete",
            "Graceful degradation when Supabase offline (flash banners; empty lists)",
            "View::share('adminBase', 'dainely-admin-panel') in AppServiceProvider",
        ],
    )
    doc.add_page_break()

    # 3. Folder Structure
    h(doc, "3. Folder Structure (Admin-relevant)", 1)
    table(
        doc,
        ["Path", "Purpose"],
        [
            ["app/Http/Controllers/Admin/", "All admin controllers + AdminController base"],
            ["app/Http/Middleware/AdminAuth.php", "Session gate for protected routes"],
            ["app/Models/Supabase/", "CMS Eloquent models (connection supabase)"],
            ["app/Models/Catalog/EducationPage.php", "Virtual education entity (no DB table)"],
            ["app/Support/SupabaseDb.php", "Availability, ping, safe run() wrapper"],
            ["app/Support/ContentCatalog.php", "Static education + blog catalog IDs"],
            ["app/Jobs/RetryFailedWebhooksJob.php", "Used by admin webhook retry + scheduler"],
            ["resources/views/layouts/admin.blade.php", "Admin shell / sidebar"],
            ["resources/views/admin/", "Module Blade templates"],
            ["routes/web.php", "dainely-admin-panel route group"],
            ["config/supabase.php", "FEATURES_SUPABASE toggle + API keys"],
            ["database/migrations/2026_07_*", "Supabase schema + RLS migration"],
        ],
    )
    doc.add_page_break()

    # 4. Auth
    h(doc, "4. Authentication & Authorization", 1)
    p(doc, "Admin uses a single shared credential pair from environment variables, compared in plain text at login. Success stores session('admin_authenticated', true).")

    h(doc, "4.1 Login Flow", 2)
    numbered(
        doc,
        [
            "GET /dainely-admin-panel/login → AdminAuthController@login",
            "If already authenticated → redirect dashboard",
            "POST login with email + password",
            "Compare to config/env ADMIN_EMAIL and ADMIN_PASSWORD (code fallbacks exist if unset)",
            "On success: session put admin_authenticated; redirect dashboard with success flash",
            "On failure: back with validation errors",
        ],
    )

    h(doc, "4.2 Logout", 2)
    p(doc, "POST /dainely-admin-panel/logout clears admin_authenticated and redirects to login.")

    h(doc, "4.3 Not Implemented", 2)
    bullets(
        doc,
        [
            "Password reset / forgot password",
            "Remember me",
            "MFA (ADMIN_MFA_ENABLED in .env.example is not enforced in code)",
            "Roles / permissions / policies (Spatie unused)",
            "Per-module authorization — any authenticated admin has full access",
            "Login rate limiting",
            "Session regeneration on login",
        ],
    )

    h(doc, "4.4 Credentials", 2)
    table(
        doc,
        ["Source", "Default if unset (in code)"],
        [
            ["ADMIN_EMAIL / config app.admin_email", "admin@dainelylab.com"],
            ["ADMIN_PASSWORD / config app.admin_password", "admin12345"],
        ],
    )
    p(doc, "Security note: passwords are compared as plain strings from env, not hashed user records. Ensure production .env overrides defaults.", italic=True)
    doc.add_page_break()

    # 5. Database
    h(doc, "5. Database Documentation (Admin CMS)", 1)
    p(doc, "All admin CRUD targets the supabase PostgreSQL connection. Default Laravel MySQL/SQLite is not used for CMS overlays.")

    h(doc, "5.1 Tables", 2)
    table(
        doc,
        ["Table", "Purpose", "Key relationships"],
        [
            ["dainely_products", "Shopify product mirror (admin read-only list)", "hasMany product_content, morph page_blocks/faqs"],
            ["product_content", "Per-locale SEO + copy overlays", "belongsTo product"],
            ["landing_pages", "Marketing landing pages + offer meta", "morph page_blocks/faqs; optional bundle"],
            ["page_blocks", "CMS blocks on product/landing/education", "morphTo blockable"],
            ["faqs", "Polymorphic FAQs", "morphTo faqable"],
            ["product_bundles", "Bundle offers", "hasMany items"],
            ["product_bundle_items", "Bundle line items", "belongsTo bundle + product"],
            ["related_content", "Internal knowledge graph links", "string type + id (no Eloquent morph)"],
            ["product_knowledge_signals", "AI Q&A signals for products", "belongsTo product"],
            ["webhook_logs", "Inbound webhook processing log", "retry columns: attempts, next_retry_at"],
            ["user_activity_log", "Dashboard recent activity", "morphTo item"],
            ["analytics_events", "Storefront analytics (not written by admin UI)", "standalone"],
            ["search_index", "FTS index (admin does not edit directly)", "morph searchable"],
            ["ai_schema_cache", "JSON-LD cache", "morph cacheable"],
            ["recommendation_rules", "Upsell rules (not exposed in admin UI)", "morph source/recommended"],
        ],
    )

    h(doc, "5.2 Virtual entities (no table)", 2)
    p(doc, "EducationPage and blog posts for related links come from ContentCatalog (PHP arrays). Morph keys store App\\Models\\Catalog\\EducationPage with catalog IDs 1–6. There is no education_pages table — eager-loading MorphTo would fail (fixed via AdminFaqController::attachFaqables).")

    h(doc, "5.3 ER Summary (textual)", 2)
    code(
        doc,
        "dainely_products 1──* product_content\n"
        "dainely_products 1──* product_knowledge_signals\n"
        "dainely_products / landing_pages / EducationPage(catalog)\n"
        "        └── morph ──* page_blocks\n"
        "        └── morph ──* faqs\n"
        "product_bundles 1──* product_bundle_items *──1 dainely_products\n"
        "related_content: source_type/id → related_type/id",
    )
    doc.add_page_break()

    # 6. Features
    h(doc, "6. Feature Documentation", 1)

    modules = [
        (
            "6.1 Dashboard",
            "CMS health overview: synced product count, landings, bundles, failed/dead webhooks, recent user activity.",
            "Menu: Dashboard · GET /dainely-admin-panel/dashboard · AdminDashboardController@index · View: admin/dashboard.blade.php",
            [
                "flashIfSupabaseOffline('CMS Dashboard')",
                "Counts via SupabaseDb::run fallbacks to 0 / empty collection",
                "Recent UserActivityLog limit 10",
            ],
            ["Dashboard loads with metrics when Supabase online", "Offline banner when ping fails", "Activity empty state"],
        ),
        (
            "6.2 Products Overlay",
            "Edit per-locale SEO/copy overlays and page blocks for Shopify-synced products. Does not create Shopify products.",
            "Menu: Products Overlay · /products · AdminProductController · Views: admin/products/*",
            [
                "edit(): firstOrCreate ProductContent for en/fr/de",
                "update(): empty strings → null (storefront falls back to Shopify/lang)",
                "Block CRUD on morph PageBlock",
                "Cache forget: product_*, json_ld_*, storefront.pdp_overlay.*",
            ],
            [
                "List synced products",
                "Save SEO title/description → live meta updates",
                "Clear SEO fields → Shopify/lang defaults return",
                "Add/edit/delete blocks",
            ],
        ),
        (
            "6.3 Landing Pages",
            "Create/edit marketing landings, publish flag, offer fields (Shopify product, bundle, CTA, discount), and page blocks.",
            "Menu: Landing Pages · /landings · AdminLandingController",
            [
                "store/update validation for slug, locale, meta, offer fields",
                "discount_code normalized uppercase on update",
                "Blocks inherit landing locale",
                "Invalidates landing cache keys",
            ],
            ["Create landing", "Edit offer fields", "Block CRUD", "Published flag"],
        ),
        (
            "6.4 Education Blocks",
            "Manage page_blocks for six static education topics from ContentCatalog (not DB pages).",
            "Menu: Education Blocks · /education · AdminEducationController",
            [
                "index lists ContentCatalog::educationPages()",
                "edit uses EducationPage::findCatalog()",
                "Blocks stored with blockable_type = EducationPage::class",
                "No education page create/delete in admin",
            ],
            ["List 6 education pages", "Invalid ID redirects", "Block add/update/delete"],
        ),
        (
            "6.5 Bundles & Offers",
            "Create bundles linked to a Shopify bundle product GID and attach component products with quantities.",
            "Menu: Bundles · /bundles · AdminBundleController",
            [
                "Unique bundle_shopify_product_id on supabase connection",
                "addItem quantity 1–10; duplicate product guard",
                "No bundle delete route",
            ],
            ["Create bundle", "Reject duplicate Shopify GID", "Add/remove items"],
        ),
        (
            "6.6 FAQs Manager",
            "Polymorphic FAQs for Product, LandingPage, or EducationPage catalog entries.",
            "Menu: FAQs · /faqs · AdminFaqController",
            [
                "attachFaqables() avoids education_pages table query",
                "store sets approved=true",
                "Optional ?locale= filter",
                "Inline Alpine edit UI",
            ],
            ["Create FAQ for each target type", "Education invalid ID rejected", "Update/delete", "List shows parent title"],
        ),
        (
            "6.7 Knowledge Signals",
            "Review AI/customer/expert Q&A signals attached to products; approve or edit.",
            "Menu: Knowledge Signals · /signals · AdminSignalController",
            [
                "Filters: locale, approved",
                "toggleApproval flips boolean",
                "speaker_type in: expert,customer,ai",
                "Clears product cache keys",
            ],
            ["Filter list", "Toggle approval", "Update Q&A"],
        ),
        (
            "6.8 Internal Links (Knowledge Graph)",
            "Create ordered relationships between product, landing_page, education, blog catalog entities.",
            "Menu: Internal Links · /related · AdminRelatedController",
            [
                "Types are strings: product|landing_page|education|blog",
                "Rejects self-links and duplicates",
                "entityExists() validates targets",
                "update only display_order",
            ],
            ["Create link", "Reject self/duplicate", "Update order", "Delete"],
        ),
        (
            "6.9 Webhook Logs",
            "Inspect webhook processing logs and manually retry failed/retryable entries.",
            "Menu: Webhook Logs · /webhooks · AdminWebhookController",
            [
                "Paginated WebhookLog",
                "retry() resets log and runs RetryFailedWebhooksJob(100) synchronously",
                "Scheduler also runs job every 5 minutes",
            ],
            ["List logs", "Retry updates status flash"],
        ),
    ]

    for title, purpose, nav, logic, tests in modules:
        h(doc, title, 2)
        p(doc, "Purpose: " + purpose)
        p(doc, nav, italic=True)
        p(doc, "Business / technical logic:", bold=True)
        bullets(doc, logic)
        p(doc, "QA checklist:", bold=True)
        bullets(doc, ["✓ " + t for t in tests])
        p(doc, "[Screenshot placeholder]", italic=True)

    doc.add_page_break()

    # 7. Routes summary already in tables - condensed
    h(doc, "7. Routes Summary", 1)
    p(doc, "All admin routes live under prefix dainely-admin-panel with middleware web. Protected routes also use admin.auth. Named routes: admin.login, admin.logout, admin.dashboard. Full method/URI list is derived from routes/web.php (41 entries including root redirect).")
    p(doc, "Root GET /dainely-admin-panel/ redirects to dashboard if authenticated, otherwise login.")
    doc.add_page_break()

    # 8 Controllers
    h(doc, "8. Controllers", 1)
    table(
        doc,
        ["Controller", "Extends", "Key methods"],
        [
            ["AdminAuthController", "Controller", "login, authenticate, logout"],
            ["AdminController", "Controller (abstract)", "flashIfSupabaseOffline"],
            ["AdminDashboardController", "AdminController", "index"],
            ["AdminProductController", "AdminController", "index, edit, update, addBlock, updateBlock, deleteBlock"],
            ["AdminLandingController", "AdminController", "index, store, edit, update, addBlock, updateBlock, deleteBlock"],
            ["AdminEducationController", "AdminController", "index, edit, addBlock, updateBlock, deleteBlock"],
            ["AdminBundleController", "AdminController", "index, store, edit, update, addItem, deleteItem"],
            ["AdminFaqController", "AdminController", "index, store, update, destroy, attachFaqables"],
            ["AdminSignalController", "AdminController", "index, toggleApproval, update"],
            ["AdminRelatedController", "AdminController", "index, store, update, destroy, entityExists"],
            ["AdminWebhookController", "AdminController", "index, retry"],
        ],
    )
    doc.add_page_break()

    # 9 Models short
    h(doc, "9. Models (Admin-used)", 1)
    p(doc, "See Database section for tables. All Supabase models set protected $connection = 'supabase'. Soft deletes are not used on CMS overlay models. Product booted() hooks queue search indexing on save/delete.")
    doc.add_page_break()

    # 10 Services
    h(doc, "10. Services & Support (Admin touchpoints)", 1)
    table(
        doc,
        ["Class", "Admin relevance"],
        [
            ["SupabaseDb", "enabled/driverLoaded/ping/available/run — every module"],
            ["ContentCatalog", "Education + blog lists for FAQs/Related/Education"],
            ["RetryFailedWebhooksJob", "Admin webhook retry + schedule"],
            ["SearchService", "Indirect via ProductContent saved hooks"],
        ],
    )
    p(doc, "Admin does not call ShopifyService, SquareService, or payment services directly.")
    doc.add_page_break()

    # 11 Validation
    h(doc, "11. Validation Rules", 1)
    p(doc, "All validation is inline in controllers. Highlights:")
    bullets(
        doc,
        [
            "Product overlay: contents.* fields nullable strings; seo_title max 255; seo_description max 1000",
            "Blocks: block_type in benefits|how-it-works|testimonials|faqs|cta|video|comparison|bundle",
            "Landing: slug/title required; published boolean; discount_code max 50",
            "Bundle: unique bundle_shopify_product_id on supabase connection; quantity 1–10",
            "FAQ: faqable_type must be Product|LandingPage|EducationPage class names",
            "Signals: speaker_type in expert|customer|ai",
            "Related: types product|landing_page|education|blog",
        ],
    )
    doc.add_page_break()

    # 12 Middleware
    h(doc, "12. Middleware", 1)
    table(
        doc,
        ["Middleware", "Purpose"],
        [
            ["web group", "Session, CSRF, cookies, bindings, TrackPageViews"],
            ["admin.auth", "Require session admin_authenticated"],
            ["TrackPageViews", "Skips dainely-admin-panel/* (no analytics writes)"],
            ["locale", "NOT applied to admin (storefront only)"],
        ],
    )
    doc.add_page_break()

    # 13-15 Jobs
    h(doc, "13–15. Scheduler, Queues, Events", 1)
    bullets(
        doc,
        [
            "Scheduler: RetryFailedWebhooksJob every 5 minutes (Console Kernel)",
            "Admin webhook retry runs the same job synchronously",
            "Admin CRUD is synchronous (no queued jobs for save operations)",
            "Hourly reviews:warm-cache is storefront-related, not admin UI",
            "No admin-specific Events/Listeners for CMS saves",
            "AnalyticsEventOccurred is storefront analytics pipeline",
        ],
    )
    doc.add_page_break()

    # 16 Notifications
    h(doc, "16. Notifications", 1)
    p(doc, "Admin panel uses session flash messages only (success/error). No email/SMS/Slack notifications are sent from Admin controllers.")
    doc.add_page_break()

    # 17 Uploads
    h(doc, "17. File Uploads", 1)
    p(doc, "Admin CMS forms do not upload images. Content is text/HTML/Markdown in textareas. Product images come from Shopify sync (featured_image URL).")
    doc.add_page_break()

    # 18 APIs
    h(doc, "18. APIs", 1)
    p(doc, "Admin UI is server-rendered Blade (no admin JSON API). Related external systems:")
    bullets(
        doc,
        [
            "Supabase PostgreSQL via PDO (primary)",
            "Inbound webhooks at /api/webhooks/* and /webhooks/* (Shopify, Judge.me, etc.) — inspected via Webhook Logs",
            "Shopify Admin/Storefront APIs used by sync jobs, not by admin forms",
        ],
    )
    doc.add_page_break()

    # 19 Config
    h(doc, "19. Configuration & Environment", 1)
    table(
        doc,
        ["Variable", "Purpose"],
        [
            ["ADMIN_EMAIL / ADMIN_PASSWORD", "Admin login credentials"],
            ["ADMIN_MFA_ENABLED", "Documented in env.example; not enforced in code"],
            ["FEATURES_SUPABASE", "Enable/disable Supabase CMS"],
            ["DB_SUPABASE_*", "PostgreSQL pooler connection"],
            ["CACHE_DRIVER", "Used for admin ping cache + overlay invalidation"],
            ["QUEUE_CONNECTION", "Affects storefront jobs; admin retry runs sync job call"],
        ],
    )
    doc.add_page_break()

    # 20 Security
    h(doc, "20. Security Review", 1)
    bullets(
        doc,
        [
            "CSRF: enabled on all admin POSTs",
            "XSS: Blade escaping by default; some content fields intentionally allow HTML when rendered on storefront",
            "SQL injection: Eloquent/query builder parameterized",
            "Mass assignment: fillable arrays on models",
            "Auth: single shared env password (plain compare) — rotate and avoid defaults in production",
            "No per-feature authorization",
            "Supabase RLS enabled for public API; Laravel owner connection continues",
            "Recommendation: hash admin credentials or move to User + hashed password; add rate limit on login; regenerate session on login",
        ],
    )
    doc.add_page_break()

    # 21 Errors
    h(doc, "21. Error Handling", 1)
    bullets(
        doc,
        [
            "SupabaseDb::run catches connection failures and returns fallbacks",
            "Controllers flash human-readable offline reasons via failureReason()",
            "Validation errors return back() with $errors",
            "Laravel exception pages when APP_DEBUG=true (e.g. historic education_pages MorphTo error)",
        ],
    )
    doc.add_page_break()

    # 22 Workflows
    h(doc, "22. Admin Workflows", 1)
    h(doc, "22.1 Product Overlay → Storefront Meta", 2)
    code(
        doc,
        "Admin logs in\n→ Products Overlay → Edit product\n→ Fill SEO Title / Description (or leave empty)\n→ Save All Locales\n→ Cache keys cleared\n→ Storefront PDP: CMS values if filled, else lang file, else Shopify",
    )
    h(doc, "22.2 Webhook Failure Recovery", 2)
    code(
        doc,
        "Inbound webhook → webhook_logs row\n→ Auto RetryFailedWebhooksJob every 5 min\n→ OR Admin Webhook Logs → Retry\n→ Job reprocesses; status updated",
    )
    doc.add_page_break()

    # 23 QA
    h(doc, "23. QA Testing Guide (Cross-cutting)", 1)
    p(doc, "Prerequisites: PHP with pdo_pgsql, FEATURES_SUPABASE=true, valid DB_SUPABASE_*, ADMIN_* set, npm build assets available.")
    bullets(
        doc,
        [
            "Auth: wrong password; correct password; logout; deep-link protected URL when logged out",
            "CSRF: POST without token → 419",
            "Offline: disable FEATURES_SUPABASE or break DB → banners, no fatal on list pages",
            "Products: SEO filled vs empty fallback on live PDP",
            "FAQs: Product + Landing + Education targets; education invalid ID",
            "Related: self-link and duplicate rejected",
            "Bundles: duplicate Shopify GID rejected",
            "Commands: php artisan admin:verify ; php artisan supabase:diagnose",
            "Regression: storefront still loads with overlays; admin path is /dainely-admin-panel not /admin",
        ],
    )
    doc.add_page_break()

    # 24 Troubleshooting
    h(doc, "24. Troubleshooting", 1)
    table(
        doc,
        ["Symptom", "Cause", "Resolution"],
        [
            ["pdo_pgsql / offline banners", "Missing PHP extension or bad DB_SUPABASE_*", "Enable pgsql; fix pooler user; config:clear"],
            ["education_pages does not exist", "MorphTo eager load on virtual EducationPage", "Use attachFaqables (current code); deploy AdminFaqController"],
            ["Overlay changes not on live", "Cache / blade used lang only (historical)", "Deploy ProductController + show.blade SEO priority; hard refresh"],
            ["False Database query failed flash", "Eager redirect()->with() fallback", "Use fn () => redirect()->with() (current)"],
            ["Login fails", "Wrong env / special chars", "Verify ADMIN_* ; no trailing spaces"],
            ["Webhook retry no effect", "Log not retryable / job errors", "Inspect webhook_logs; check schedule:run"],
        ],
    )
    doc.add_page_break()

    # 25 Deployment
    h(doc, "25. Deployment Notes (Admin)", 1)
    numbered(
        doc,
        [
            "Deploy PHP/Blade routes including dainely-admin-panel prefix",
            "Ensure pdo_pgsql on host PHP version",
            "Set ADMIN_EMAIL, ADMIN_PASSWORD, DB_SUPABASE_*, FEATURES_SUPABASE=true",
            "Run migrations including RLS enablement on supabase connection",
            "php artisan config:clear && route:clear && view:clear && cache:clear",
            "Confirm cron runs schedule:run for webhook retries",
            "Verify: /dainely-admin-panel/login and admin:verify",
        ],
    )
    doc.add_page_break()

    # 26-27 Debt
    h(doc, "26. Known Technical Debt", 1)
    bullets(
        doc,
        [
            "Plain-text admin password in env with insecure code defaults",
            "No RBAC despite Spatie package",
            "Large Blade product landing templates vs CMS overlays (dual content sources)",
            "Hardcoded sidebar URLs vs $adminBase inconsistency",
            "Docs may still mention /admin paths",
            "Dashboard loads pending webhook count but UI may not show it",
        ],
    )

    h(doc, "27. Suggested Refactoring", 1)
    bullets(
        doc,
        [
            "Move admin credentials to hashed User model or Sanctum token",
            "Form Request classes + Policies if multi-admin roles needed",
            "Central AdminUrl helper for path prefix",
            "Eager-load strategies unified for morph parents",
            "Extract block CRUD into shared service/trait",
        ],
    )
    doc.add_page_break()

    # 28 Onboarding
    h(doc, "28. Developer Onboarding (Admin)", 1)
    numbered(
        doc,
        [
            "Clone repo; composer install; copy .env.example → .env; php artisan key:generate",
            "Enable pdo_pgsql; set DB_SUPABASE_* (pooler) and FEATURES_SUPABASE=true",
            "Set ADMIN_EMAIL / ADMIN_PASSWORD",
            "php artisan migrate (including supabase migrations)",
            "npm ci && npm run build",
            "php artisan serve → open /dainely-admin-panel/login",
            "Useful: php artisan admin:verify ; supabase:diagnose ; config:clear ; cache:clear",
        ],
    )
    doc.add_page_break()

    # 29 Matrix
    h(doc, "29. Feature Dependency Matrix", 1)
    table(
        doc,
        ["Feature", "DB", "Services", "Queue", "Permissions"],
        [
            ["Dashboard", "multiple supabase tables", "SupabaseDb", "—", "session only"],
            ["Products Overlay", "dainely_products, product_content, page_blocks", "SupabaseDb + cache", "search index job on content save", "session only"],
            ["Landings", "landing_pages, page_blocks", "SupabaseDb + cache", "—", "session only"],
            ["Education", "page_blocks + ContentCatalog", "SupabaseDb", "—", "session only"],
            ["Bundles", "product_bundles, items", "SupabaseDb", "—", "session only"],
            ["FAQs", "faqs", "SupabaseDb + ContentCatalog", "—", "session only"],
            ["Signals", "product_knowledge_signals", "SupabaseDb + cache", "—", "session only"],
            ["Related", "related_content", "SupabaseDb + ContentCatalog", "—", "session only"],
            ["Webhooks UI", "webhook_logs", "RetryFailedWebhooksJob", "schedule 5 min", "session only"],
        ],
    )
    doc.add_page_break()

    # 30 Appendix
    h(doc, "30. Appendix — Glossary", 1)
    table(
        doc,
        ["Term", "Meaning"],
        [
            ["Overlay", "Supabase content that overrides Shopify/lang copy or SEO"],
            ["adminBase", "URL prefix dainely-admin-panel shared to Blade"],
            ["ContentCatalog", "Static PHP list of education/blog entities"],
            ["EducationPage", "Virtual model; no education_pages table"],
            ["SupabaseDb::run", "Execute callback or return fallback on connection failure"],
            ["Morph blockable/faqable", "Polymorphic parent for blocks/FAQs"],
            ["RLS", "Row Level Security — API locked; Laravel PDO owner still works"],
        ],
    )

    p(doc, "End of Admin Panel documentation. Generated from application source under Dainely-Premium-Wellness. No functionality invented beyond the codebase.", italic=True)

    doc.save(OUT)
    print(f"Wrote: {OUT}")
    print(f"Size: {OUT.stat().st_size} bytes")


if __name__ == "__main__":
    build()
