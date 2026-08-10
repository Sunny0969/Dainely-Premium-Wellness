#!/usr/bin/env python3
"""Merge docs/*.md into a single Word document."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

DOCS_DIR = Path(__file__).resolve().parents[1] / "docs"
OUT_PATH = Path(__file__).resolve().parents[1] / "Dainely-Premium-Wellness-Technical-Documentation.docx"

# Ordered sections matching the documentation deliverables
DOC_ORDER = [
    "README.md",
    "Folder-Structure.md",
    "Features.md",
    "Architecture.md",
    "Database.md",
    "Routes.md",
    "Models.md",
    "Services.md",
    "API.md",
    "Queues.md",
    "Integrations.md",
    "Environment.md",
    "Configuration.md",
    "Permissions.md",
    "Frontend.md",
    "Deployment.md",
    "Testing.md",
    "Troubleshooting.md",
    "Coding-Standards.md",
    "Developer-Guide.md",
    "AI-Guide.md",
    "CodeMap.md",
    "CHANGELOG.md",
]

SECTION_TITLES = {
    "README.md": "1. Project Overview",
    "Folder-Structure.md": "2. Folder Structure Guide",
    "Features.md": "3. Feature Catalog",
    "Architecture.md": "4. Application Architecture",
    "Database.md": "5. Database Documentation",
    "Routes.md": "6. Route Documentation",
    "Models.md": "7. Models",
    "Services.md": "8. Services",
    "API.md": "9. API Documentation",
    "Queues.md": "10. Queue & Jobs",
    "Integrations.md": "11. Integrations",
    "Environment.md": "12. Environment Variables",
    "Configuration.md": "13. Configuration Guide",
    "Permissions.md": "14. Permissions",
    "Frontend.md": "15. Frontend",
    "Deployment.md": "16. Deployment",
    "Testing.md": "17. Testing",
    "Troubleshooting.md": "18. Troubleshooting",
    "Coding-Standards.md": "19. Coding Standards",
    "Developer-Guide.md": "20. Developer Quick Start",
    "AI-Guide.md": "21. AI Assistant Guide",
    "CodeMap.md": "22. Code Map",
    "CHANGELOG.md": "23. Change Log",
}


def set_run_font(run, name: str = "Calibri", size: int | None = None, bold: bool = False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    run.bold = bold


def add_heading_styled(doc: Document, text: str, level: int):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return h


def strip_md_links(text: str) -> str:
    text = re.sub(r"!\[([^\]]*)\]\([^)]+\)", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    return text


def inline_runs(paragraph, text: str):
    """Parse **bold**, `code`, and plain text into runs."""
    text = strip_md_links(text)
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*)")
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            run = paragraph.add_run(text[pos : m.start()])
            set_run_font(run, size=11)
        token = m.group(0)
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=11, bold=True)
        elif token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, name="Consolas", size=10)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        elif token.startswith("*") and token.endswith("*"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size=11)
            run.italic = True
        pos = m.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, size=11)


def add_paragraph_md(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    inline_runs(p, text)
    return p


def add_code_block(doc: Document, code: str):
    for line in code.splitlines() or [""]:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Inches(0.25)
        run = p.add_run(line if line else " ")
        set_run_font(run, name="Consolas", size=9)
        run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
        # Light gray background via shading
        shd = p._element.get_or_add_pPr()
        shading = shd.makeelement(
            qn("w:shd"),
            {
                qn("w:val"): "clear",
                qn("w:color"): "auto",
                qn("w:fill"): "F5F5F5",
            },
        )
        shd.append(shading)
    doc.add_paragraph()


def add_table(doc: Document, headers: list[str], rows: list[list[str]]):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        run = p.add_run(strip_md_links(h).strip())
        set_run_font(run, size=10, bold=True)
        # Header shading
        tc = hdr[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shading = tcPr.makeelement(
            qn("w:shd"),
            {
                qn("w:val"): "clear",
                qn("w:color"): "auto",
                qn("w:fill"): "E8EEF4",
            },
        )
        tcPr.append(shading)
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx in range(len(headers)):
            val = row[c_idx] if c_idx < len(row) else ""
            cells[c_idx].text = ""
            p = cells[c_idx].paragraphs[0]
            inline_runs(p, val.strip())
            for run in p.runs:
                run.font.size = Pt(9)
    doc.add_paragraph()


def parse_table_block(lines: list[str], start: int) -> tuple[list[str], list[list[str]], int]:
    """Parse a markdown table starting at start. Returns headers, rows, next index."""
    header_line = lines[start].strip()
    headers = [c.strip() for c in header_line.strip("|").split("|")]
    i = start + 1
    if i < len(lines) and re.match(r"^\|?\s*:?-+:?\s*(\|\s*:?-+:?\s*)+\|?$", lines[i].strip()):
        i += 1
    rows = []
    while i < len(lines) and lines[i].strip().startswith("|"):
        cols = [c.strip() for c in lines[i].strip().strip("|").split("|")]
        # Pad/truncate to header count
        while len(cols) < len(headers):
            cols.append("")
        rows.append(cols[: len(headers)])
        i += 1
    return headers, rows, i


def process_markdown(doc: Document, content: str, skip_first_h1: bool = True):
    lines = content.replace("\r\n", "\n").split("\n")
    i = 0
    first_h1_skipped = False
    in_code = False
    code_lang = ""
    code_buf: list[str] = []

    while i < len(lines):
        line = lines[i]

        # Fenced code blocks
        fence = re.match(r"^```(\w*)\s*$", line.strip())
        if fence:
            if not in_code:
                in_code = True
                code_lang = fence.group(1) or ""
                code_buf = []
            else:
                code = "\n".join(code_buf)
                if code_lang == "mermaid":
                    p = doc.add_paragraph()
                    run = p.add_run("[Mermaid diagram — see docs/*.md or render in a Mermaid viewer]")
                    set_run_font(run, size=10)
                    run.italic = True
                    add_code_block(doc, code)
                else:
                    add_code_block(doc, code)
                in_code = False
                code_buf = []
            i += 1
            continue

        if in_code:
            code_buf.append(line)
            i += 1
            continue

        # Horizontal rule
        if re.match(r"^---+\s*$", line.strip()):
            doc.add_paragraph("─" * 40)
            i += 1
            continue

        # Empty line
        if not line.strip():
            i += 1
            continue

        # Headings
        hm = re.match(r"^(#{1,6})\s+(.+)$", line)
        if hm:
            level = len(hm.group(1))
            text = strip_md_links(hm.group(2).strip())
            if level == 1 and skip_first_h1 and not first_h1_skipped:
                first_h1_skipped = True
                i += 1
                continue
            # Cap heading levels for Word (1–3 under section H1)
            word_level = min(level + 1, 4) if skip_first_h1 else min(level, 4)
            add_heading_styled(doc, text, word_level)
            i += 1
            continue

        # Tables
        if line.strip().startswith("|") and i + 1 < len(lines):
            nxt = lines[i + 1].strip()
            if re.match(r"^\|?\s*:?-+:?\s*(\|\s*:?-+:?\s*)+\|?$", nxt) or nxt.startswith("|"):
                headers, rows, i = parse_table_block(lines, i)
                add_table(doc, headers, rows)
                continue

        # Unordered list
        lm = re.match(r"^(\s*)([-*+])\s+(.+)$", line)
        if lm:
            text = lm.group(3)
            p = doc.add_paragraph(style="List Bullet")
            inline_runs(p, text)
            i += 1
            continue

        # Ordered list
        om = re.match(r"^(\s*)\d+\.\s+(.+)$", line)
        if om:
            text = om.group(2)
            p = doc.add_paragraph(style="List Number")
            inline_runs(p, text)
            i += 1
            continue

        # Blockquote
        if line.strip().startswith(">"):
            text = re.sub(r"^>\s?", "", line.strip())
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            run = p.add_run(strip_md_links(text))
            set_run_font(run, size=11)
            run.italic = True
            i += 1
            continue

        # Normal paragraph (join soft-wrapped? keep line-based for simplicity)
        add_paragraph_md(doc, line.strip())
        i += 1


def build():
    doc = Document()

    # Narrow margins slightly for denser docs
    for section in doc.sections:
        section.top_margin = Inches(0.85)
        section.bottom_margin = Inches(0.85)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    # Cover
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(72)
    run = title.add_run("Dainely Premium Wellness")
    set_run_font(run, size=28, bold=True)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Complete Technical Documentation")
    set_run_font(run, size=18)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x55)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_before = Pt(24)
    run = meta.add_run(
        "Laravel 10 storefront & Phase 2 Admin CMS\n"
        "For developers and AI coding assistants\n"
        "Generated from /docs — secrets omitted"
    )
    set_run_font(run, size=11)

    toc_note = doc.add_paragraph()
    toc_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    toc_note.paragraph_format.space_before = Pt(36)
    run = toc_note.add_run(
        "Contents: Overview · Structure · Features · Architecture · Database · "
        "Routes · Models · Services · API · Queues · Integrations · Environment · "
        "Configuration · Permissions · Frontend · Deployment · Testing · "
        "Troubleshooting · Standards · Developer Guide · AI Guide · Code Map · Changelog"
    )
    set_run_font(run, size=9)
    run.italic = True

    doc.add_page_break()

    missing = []
    for filename in DOC_ORDER:
        path = DOCS_DIR / filename
        if not path.exists():
            missing.append(filename)
            continue

        section_title = SECTION_TITLES.get(filename, filename)
        add_heading_styled(doc, section_title, 1)

        content = path.read_text(encoding="utf-8")
        process_markdown(doc, content, skip_first_h1=True)
        doc.add_page_break()

    if missing:
        add_heading_styled(doc, "Missing source files", 1)
        for m in missing:
            doc.add_paragraph(m, style="List Bullet")

    # Remove trailing empty page break if last
    # (python-docx keeps the last page break; acceptable)

    doc.save(OUT_PATH)
    print(f"Wrote: {OUT_PATH}")
    print(f"Sections: {len(DOC_ORDER) - len(missing)}/{len(DOC_ORDER)}")
    if missing:
        print("Missing:", ", ".join(missing))


if __name__ == "__main__":
    build()
